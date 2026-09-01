# SPDX-License-Identifier: Apache-2.0
# Copyright (c) 2026 Amir Farhadi

import os
import shutil
import subprocess
import time
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader

def convert_audiobook_markdown_to_docx_and_pdf(md_path: Path):
    docx_path = md_path.with_suffix(".docx")
    pdf_path = md_path.with_suffix(".pdf")
    html_path = md_path.with_suffix(".html")

    with open(md_path, "r", encoding="utf-8") as f:
        raw_lines = f.readlines()

    doc = Document()

    # Configure Margins and ensure empty headers/footers on all pages
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    html_body = []

    in_code_block = False
    code_lines = []
    current_list_type = None  # 'ul' or 'ol'
    doc_title = md_path.stem.replace("_", " ")

    def close_current_list():
        nonlocal current_list_type
        if current_list_type == 'ul':
            html_body.append("</ul>")
            current_list_type = None
        elif current_list_type == 'ol':
            html_body.append("</ol>")
            current_list_type = None

    for raw_line in raw_lines:
        line = raw_line.rstrip("\r\n")

        # Skip license and metadata comments
        if line.startswith("<!--") or line.endswith("-->"):
            continue
        if line.strip() == "---":
            close_current_list()
            doc.add_paragraph("")
            html_body.append("<hr/>")
            continue

        if line.startswith("```"):
            close_current_list()
            if in_code_block:
                in_code_block = False
                code_text = "\n".join(code_lines)

                # Word
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

                # HTML
                escaped = code_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                html_body.append(f"<pre><code>{escaped}</code></pre>")
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        stripped = line.strip()
        if not stripped:
            close_current_list()
            continue

        # Headings
        if stripped.startswith("# "):
            close_current_list()
            title_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(title_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            html_body.append(f"<h1>{title_text}</h1>")

        elif stripped.startswith("## "):
            close_current_list()
            h2_text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(h2_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            html_body.append(f"<h2>{h2_text}</h2>")

        elif stripped.startswith("### "):
            close_current_list()
            h3_text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(h3_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            html_body.append(f"<h3>{h3_text}</h3>")

        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:].strip().replace("**", "").replace("`", "")
            if current_list_type != 'ul':
                close_current_list()
                html_body.append("<ul>")
                current_list_type = 'ul'
            html_body.append(f"<li>{bullet_text}</li>")

            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(bullet_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ('.', ')'):
            item_text = stripped[2:].strip().replace("**", "").replace("`", "")
            if current_list_type != 'ol':
                close_current_list()
                html_body.append("<ol>")
                current_list_type = 'ol'
            html_body.append(f"<li>{item_text}</li>")

            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(item_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

        else:
            close_current_list()
            clean_text = stripped.replace("**", "").replace("`", "")
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(clean_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            html_body.append(f"<p>{clean_text}</p>")

    close_current_list()

    doc.save(str(docx_path))
    print(f"DOCX updated: {docx_path}")

    # Build clean HTML without running headers or page noise
    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{doc_title}</title>
<style>
    @page {{
        size: letter;
        margin: 0.9in 1in 0.9in 1in;
    }}
    body {{
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
        font-size: 11pt;
        line-height: 1.5;
        color: #1e293b;
        background: #ffffff;
        margin: 0;
        padding: 0;
    }}
    h1 {{
        font-size: 18pt;
        color: #0f172a;
        margin-top: 1.4em;
        margin-bottom: 0.3em;
        page-break-after: avoid;
    }}
    h2 {{
        font-size: 14pt;
        color: #1e293b;
        margin-top: 1.2em;
        margin-bottom: 0.25em;
        border-bottom: 1px solid #e2e8f0;
        padding-bottom: 0.2em;
        page-break-after: avoid;
    }}
    h3 {{
        font-size: 12pt;
        color: #334155;
        margin-top: 1em;
        margin-bottom: 0.2em;
        page-break-after: avoid;
    }}
    p {{
        margin-top: 0;
        margin-bottom: 0.7em;
        text-align: justify;
    }}
    pre {{
        background: #f8fafc;
        border: 1px solid #e2e8f0;
        border-radius: 4px;
        padding: 0.8em;
        font-family: Consolas, "Courier New", monospace;
        font-size: 9.5pt;
        line-height: 1.4;
        overflow-x: auto;
        page-break-inside: avoid;
    }}
    code {{
        font-family: Consolas, "Courier New", monospace;
    }}
    ul, ol {{
        margin-top: 0;
        margin-bottom: 0.7em;
        padding-left: 1.6em;
    }}
    li {{
        margin-bottom: 0.3em;
    }}
    hr {{
        border: 0;
        border-top: 1px solid #cbd5e1;
        margin: 1.4em 0;
    }}
</style>
</head>
<body>
{''.join(html_body)}
</body>
</html>"""

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)

    edge_exe = shutil.which("msedge")
    if not edge_exe:
        for env_var in ["ProgramFiles(x86)", "ProgramFiles", "LOCALAPPDATA"]:
            base_dir = os.environ.get(env_var)
            if base_dir:
                cand = Path(base_dir) / "Microsoft" / "Edge" / "Application" / "msedge.exe"
                if cand.exists():
                    edge_exe = str(cand)
                    break

    if not edge_exe:
        raise RuntimeError("Microsoft Edge executable could not be found for PDF generation.")

    file_url = html_path.resolve().as_uri()

    cmd = [
        edge_exe,
        "--headless",
        "--disable-gpu",
        "--allow-file-access-from-files",
        "--no-pdf-header-footer",
        f"--print-to-pdf={pdf_path}",
        file_url
    ]
    res = subprocess.run(cmd, capture_output=True, text=True)
    time.sleep(1.0)

    if pdf_path.exists():
        reader = PdfReader(str(pdf_path))
        page_count = len(reader.pages)
        sample = reader.pages[0].extract_text()
        if "ERR_FILE_NOT_FOUND" in sample:
            raise RuntimeError(f"PDF generation failed with ERR_FILE_NOT_FOUND for {md_path}")
        print(f"PDF successfully verified ({page_count} pages): {pdf_path}")
    else:
        raise RuntimeError(f"Edge PDF generation failed. Stderr: {res.stderr}")

    if html_path.exists():
        html_path.unlink()

def build_all_audiobooks():
    repo_root = Path(__file__).resolve().parent.parent
    guides_dir = repo_root / "docs" / "guides"

    md_files = sorted(list(guides_dir.rglob("AUDIOBOOK_*.md")))
    for md_file in md_files:
        print(f"Processing audiobook guide: {md_file.name}")
        convert_audiobook_markdown_to_docx_and_pdf(md_file)

if __name__ == "__main__":
    build_all_audiobooks()
